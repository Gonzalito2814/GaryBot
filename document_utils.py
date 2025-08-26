from pathlib import Path
from datetime import datetime

try:
    from docx import Document
except Exception:
    Document = None

def _sanitize_filename(name: str) -> str:
    
    bad = '<>:"/\\|?*'
    for ch in bad:
        name = name.replace(ch, "_")
    return name.strip()

def save_response_to_docx(character_name: str, question: str, answer: str, out_dir: str = "data/out") -> str:
    """
    Crea un .docx con la pregunta y respuesta del personaje.
    Requiere: pip install python-docx
    """
    if Document is None:
        raise RuntimeError("Falta python-docx. Instala con: pip install python-docx")

    Path(out_dir).mkdir(parents=True, exist_ok=True)
    safe_name = _sanitize_filename(character_name)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = Path(out_dir) / f"{safe_name}_respuesta_{ts}.docx"

    doc = Document()
    doc.add_heading(f"Respuesta de {character_name}", 0)
    doc.add_paragraph(f"Fecha/Hora: {datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph("")  

    doc.add_paragraph("Pregunta:")
    doc.add_paragraph(question)
    doc.add_paragraph("")  

    doc.add_paragraph("Respuesta:")
    for line in answer.splitlines():
        doc.add_paragraph(line)

    doc.save(file_path)
    return str(file_path)

def save_response_to_pdf(character_name: str, question: str, answer: str, out_dir: str = "data/out") -> str:
    """
    Crea un PDF simple con la pregunta y respuesta.
    Requiere: pip install reportlab
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
    except Exception:
        raise RuntimeError("Falta reportlab. Instala con: pip install reportlab")

    Path(out_dir).mkdir(parents=True, exist_ok=True)
    safe_name = _sanitize_filename(character_name)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = Path(out_dir) / f"{safe_name}_respuesta_{ts}.pdf"

    c = canvas.Canvas(str(file_path), pagesize=A4)
    width, height = A4

    x = 2 * cm
    y = height - 2 * cm
    line_height = 0.6 * cm

    def writeln(text: str):
        nonlocal y
        c.drawString(x, y, text)
        y -= line_height
        if y < 2 * cm:
            c.showPage()
            y = height - 2 * cm

    writeln(f"Respuesta de {character_name}")
    writeln(f"Fecha/Hora: {datetime.now().isoformat(timespec='seconds')}")
    writeln("")
    writeln("Pregunta:")
    for line in question.splitlines():
        writeln(line)
    writeln("")
    writeln("Respuesta:")
    for line in answer.splitlines():
        writeln(line)

    c.save()
    return str(file_path)
